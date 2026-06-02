"""Generate documentation.docx for the Project Documentation Generator."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def add_code(text):
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x88, 0x35)
    p.paragraph_format.left_indent = Inches(0.4)
    return p


def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            if isinstance(cell_text, tuple):
                text, mono = cell_text
                run = row[i].paragraphs[0].add_run(text)
                if mono:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            else:
                row[i].text = cell_text
    doc.add_paragraph()
    return t


# ==============================================================
# TITLE PAGE
# ==============================================================
title = doc.add_heading("Project Documentation Generator", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Hackathon Task 1 — Technical Documentation")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

date_p = doc.add_paragraph("June 1, 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ==============================================================
# 1. PROJECT OVERVIEW
# ==============================================================
add_heading("1. Project Overview")
add_para(
    "The Project Documentation Generator is a Python-based HTTP API service that automatically "
    "synthesizes structured project documentation from a set of heterogeneous input files "
    "(.txt, .md, .json, .py, .java). It accepts raw project materials via a REST endpoint, "
    "processes them through a two-stage LLM pipeline, and returns a validated, structured JSON "
    "document together with execution metadata and a full step-by-step trace."
)
add_para(
    "The system is designed to be hallucination-resistant: all facts are extracted with source "
    "attribution, synthesis is strictly fact-bound, and empty fields are returned instead of "
    "fabricated content. Multiple LLM backends are supported (Anthropic Claude, Google Gemini, "
    "OpenRouter) and a deterministic offline mock (FakeLLM) is provided for testing and CI."
)

# ==============================================================
# 2. ARCHITECTURE
# ==============================================================
add_heading("2. Architecture")
add_para(
    "The system is built around a two-stage pipeline that separates fact extraction from "
    "document synthesis, preventing cross-contamination and reducing hallucinations."
)

add_heading("2.1  High-Level Data Flow", 2)
for line in [
    "Input Files",
    "    |",
    "    v",
    "ingest.py  -  Normalize, detect type, truncate large files (> 20,000 chars)",
    "    |",
    "    v",
    "pipeline.py / _extract()  -  Cache lookup -> LLM extraction (facts + conflicts)",
    "    |",
    "    v",
    "pipeline.py / _synthesize()  -  LLM synthesis strictly from facts",
    "    |",
    "    v",
    "schema.py / coerce_document()  -  Schema validation & normalization",
    "    |",
    "    v",
    "{ document, metadata, trace }  -  Final JSON response",
]:
    add_code(line)

doc.add_paragraph()

add_heading("2.2  Module Breakdown", 2)
add_table(
    ["Module", "Responsibility"],
    [
        (("app/main.py", True), "FastAPI application. Exposes POST /generate_document, POST /generate_document_multipart, GET/POST /config, GET /health, and GET / (HTML dashboard)."),
        (("app/pipeline.py", True), "Orchestrates the full pipeline: ingest -> extraction -> synthesis. Manages fact caching (.fact_cache.json), character-budget batching (45,000 chars), and trace recording."),
        (("app/ingest.py", True), "Receives raw file dicts, normalizes content per type (pretty-prints JSON, tags code language), truncates oversized files, and prepares XML-tagged blocks for prompts."),
        (("app/llm.py", True), "Provides LLMBase and four implementations: AnthropicLLM, GeminiLLM, OpenRouterLLM, FakeLLM. Tracks call count and cumulative token usage. make_llm() selects the backend automatically."),
        (("app/prompts.py", True), "Defines system prompts and user message builders for both pipeline stages. All prompts enforce JSON-only output and strict source attribution."),
        (("app/schema.py", True), "Pydantic models for request/response. coerce_document() normalizes any LLM output to the strict Document schema: type-coerces fields, deduplicates list items, and fills missing fields with empty defaults."),
        (("app/jsonutil.py", True), "Robust JSON extraction from raw LLM text: finds the first balanced {...} block, strips Markdown fences, and parses. Handles partial or noisy model output."),
        (("run_cli.py", True), "CLI entry point: loads all supported files from a directory and runs the full pipeline, printing the result as formatted JSON to stdout."),
    ],
)

# ==============================================================
# 3. API REFERENCE
# ==============================================================
add_heading("3. API Reference")

add_heading("3.1  POST /generate_document", 2)
add_para("Generates a structured project document from an array of file objects.")
add_para("Request body (JSON):", bold=True)
for line in [
    '{',
    '  "files": [',
    '    { "name": "notes.md",    "content": "Project goal: ..." },',
    '    { "name": "budget.json", "content": "{\\"total\\": 150000}" }',
    '  ],',
    '  "model":              "<optional model id>",',
    '  "openrouter_api_key": "<optional>",',
    '  "anthropic_api_key":  "<optional>",',
    '  "gemini_api_key":     "<optional>"',
    '}',
]:
    add_code(line)

add_para("Response (JSON):", bold=True)
for line in [
    '{',
    '  "document": {',
    '    "project_overview":   "",',
    '    "goals":              [],',
    '    "requirements":       [],',
    '    "technical_solution": "",',
    '    "architecture":       "",',
    '    "team":               [],',
    '    "timeline":           "",',
    '    "budget":             "",',
    '    "risks":              []',
    '  },',
    '  "metadata": {',
    '    "model_name":   "claude-sonnet-4-20250514",',
    '    "llm_calls":    2,',
    '    "total_tokens": 4800,',
    '    "duration_ms":  3200',
    '  },',
    '  "trace": { "steps": [ ... ] }',
    '}',
]:
    add_code(line)

add_heading("3.2  POST /generate_document_multipart", 2)
add_para(
    "Same semantics as /generate_document but accepts multipart/form-data. "
    "Files are uploaded as binary parts; model and API keys are passed as form fields. "
    "Useful for direct browser or curl-based file uploads."
)

add_heading("3.3  GET /config  |  POST /config", 2)
add_para(
    "Reads or writes the persistent configuration stored in config.json. "
    "Supported fields: default_model, openrouter_api_key, anthropic_api_key, gemini_api_key. "
    "Keys set here act as fallbacks when no key is supplied per-request."
)

add_heading("3.4  GET /health", 2)
add_code('{ "status": "ok" }')
add_para("Liveness probe. Returns HTTP 200 when the server is running.")
doc.add_paragraph()

add_heading("3.5  GET /", 2)
add_para("Serves the HTML dashboard (app/dashboard.html) for browser-based usage.")

# ==============================================================
# 4. PIPELINE DETAILS
# ==============================================================
add_heading("4. Pipeline Details")

add_heading("4.1  Stage 1 — Extraction", 2)
add_para(
    "The LLM is instructed as a careful analyst. It reads all input files (wrapped in <file> XML "
    "tags with name and type attributes) and returns only facts that are explicitly present in the "
    "text. Each fact carries a category and a source filename. Contradictory statements between "
    "files are flagged as conflicts rather than silently resolved. Allowed categories: "
    "project_overview, goals, requirements, technical_solution, architecture, team, timeline, "
    "budget, risks."
)
add_para("Extraction prompt output schema:", bold=True)
for line in [
    '{',
    '  "facts": [',
    '    { "category": "goals",',
    '      "statement": "Reduce delivery times by 25%",',
    '      "source": "project_brief.md" }',
    '  ],',
    '  "conflicts": [',
    '    { "description": "Timeline: 8 weeks vs 10 weeks",',
    '      "sources": ["project_brief.md", "meeting_notes.md"] }',
    '  ]',
    '}',
]:
    add_code(line)

add_heading("4.2  Stage 2 — Synthesis", 2)
add_para(
    "The LLM is instructed as a technical writer. It receives the JSON list of facts and constructs "
    "document fields strictly from those facts. Every claim must cite its source file in square "
    "brackets, e.g. [project_brief.md]. Fields with no supporting facts are left empty. "
    "Conflicts are resolved with a rationale recorded in resolution_notes, which surfaces in the trace."
)

add_heading("4.3  Schema Normalization (coerce_document)", 2)
add_para("After synthesis, coerce_document() enforces the Document schema regardless of model output quality:")
add_bullet("Unknown keys are dropped.")
add_bullet("String fields that arrive as lists are joined with spaces.")
add_bullet("List fields that arrive as strings are wrapped in a single-item list.")
add_bullet("None values become empty string or empty list.")
add_bullet("Duplicate list items are removed (case-insensitive, order preserved).")

add_heading("4.4  Fact Caching", 2)
add_para(
    "Extraction results are cached in .fact_cache.json keyed by filename and MD5 hash of content. "
    "On subsequent calls, unchanged files skip the LLM extraction step, saving tokens and latency. "
    "Only files whose content hash has changed trigger a new extraction call."
)

add_heading("4.5  Batching", 2)
add_para(
    "If the total character count of input files exceeds 45,000 characters, the files are split "
    "into batches that each stay within this budget. Each batch produces one extraction LLM call. "
    "Synthesis always uses a single call regardless of batch count."
)

add_heading("4.6  Typical LLM Call Count", 2)
add_table(
    ["Scenario", "Extraction calls", "Synthesis calls"],
    [
        ("All files fit in one batch", "1", "1"),
        ("Files need N batches", "N", "1"),
        ("All files cached (no changes)", "0", "1"),
        ("No input files", "0", "0"),
    ],
)

# ==============================================================
# 5. LLM BACKENDS
# ==============================================================
add_heading("5. LLM Backends")
add_para(
    "The make_llm() factory in llm.py selects the backend automatically based on available keys "
    "and the model identifier. Selection priority:"
)
add_bullet('1. DOCGEN_FAKE=1 env var -> FakeLLM (no network, deterministic)')
add_bullet('2. OpenRouter key present -> OpenRouterLLM')
add_bullet('3. Model starts with "gemini" / contains "flash" or "pro" + Gemini key -> GeminiLLM')
add_bullet('4. Anthropic key present -> AnthropicLLM')
add_bullet('5. Gemini key present -> GeminiLLM (fallback)')
add_bullet('6. No keys -> FakeLLM')
doc.add_paragraph()

add_table(
    ["Class", "Provider", "Required key / env var"],
    [
        (("AnthropicLLM", True), "Anthropic (Claude)", "ANTHROPIC_API_KEY"),
        (("GeminiLLM", True), "Google Gemini (REST API)", "GEMINI_API_KEY or GOOGLE_API_KEY"),
        (("OpenRouterLLM", True), "OpenRouter (any model)", "OPENROUTER_API_KEY"),
        (("FakeLLM", True), "Offline mock (no network)", "DOCGEN_FAKE=1 (or no keys found)"),
    ],
)
add_para(
    "All live backends use temperature=0 for maximum determinism. "
    "Token counts and call counts are tracked on the LLMBase instance and reported in the "
    "response metadata."
)

# ==============================================================
# 6. INSTALLATION & RUNNING
# ==============================================================
add_heading("6. Installation & Running")

add_heading("6.1  Prerequisites", 2)
add_bullet("Python 3.11+")
add_bullet("pip")
add_bullet("An API key for at least one supported LLM provider (or use DOCGEN_FAKE=1 for offline mode)")

add_heading("6.2  Install Dependencies", 2)
for line in [
    "python -m venv .venv",
    "# Windows:  .venv\\Scripts\\activate",
    "# Linux/Mac: source .venv/bin/activate",
    "pip install -r requirements.txt",
]:
    add_code(line)
doc.add_paragraph()

add_heading("6.3  Configure API Keys", 2)
add_para("Option A — environment variables (Windows):")
for line in [
    "set ANTHROPIC_API_KEY=sk-...",
    "set DOCGEN_MODEL=claude-sonnet-4-20250514",
    "",
    "rem  -- or Google Gemini --",
    "set GEMINI_API_KEY=AIzaSy...",
    "set DOCGEN_MODEL=gemini-2.5-flash",
    "",
    "rem  -- or OpenRouter --",
    "set OPENROUTER_API_KEY=sk-or-v1-...",
]:
    add_code(line)
doc.add_paragraph()
add_para("Option B — edit config.json in the project root:")
for line in [
    '{',
    '  "default_model":      "openai/gpt-oss-120b:free",',
    '  "openrouter_api_key": "sk-or-v1-...",',
    '  "anthropic_api_key":  null,',
    '  "gemini_api_key":     null',
    '}',
]:
    add_code(line)
doc.add_paragraph()

add_heading("6.4  Start the Server", 2)
add_para("Windows (double-click):")
add_code("start_server.bat")
add_para("Terminal:")
add_code("python -m uvicorn app.main:app --host 0.0.0.0 --port 8000")
add_para("API: http://127.0.0.1:8000/generate_document   |   Dashboard: http://127.0.0.1:8000/")
doc.add_paragraph()

add_heading("6.5  CLI Usage", 2)
for line in [
    "# Live run (requires an API key)",
    "python run_cli.py sample_data",
    "",
    "# Offline simulation (no key needed)",
    "set DOCGEN_FAKE=1 && python run_cli.py sample_data",
    "",
    "# Windows bat shortcuts",
    "run_cli_simulation.bat           # offline over sample_data",
    "run_cli_live.bat <directory>     # live generation over any directory",
]:
    add_code(line)
doc.add_paragraph()

add_heading("6.6  Run Tests", 2)
add_code("python tests/test_pipeline.py        # offline, no server needed")
add_code("python tests/test_api_client.py      # requires running server")
doc.add_paragraph()

# ==============================================================
# 7. CONFIGURATION REFERENCE
# ==============================================================
add_heading("7. Configuration Reference")
add_table(
    ["Key / Env var", "Default", "Description"],
    [
        (("DOCGEN_MODEL / default_model", True), "claude-sonnet-4-20250514", "LLM model identifier. Set to a Gemini or OpenRouter model id as needed."),
        (("ANTHROPIC_API_KEY / anthropic_api_key", True), "null", "Anthropic API key."),
        (("GEMINI_API_KEY / gemini_api_key", True), "null", "Google Gemini API key. GOOGLE_API_KEY is also accepted."),
        (("OPENROUTER_API_KEY / openrouter_api_key", True), "null", "OpenRouter API key."),
        (("DOCGEN_FAKE", True), "(unset)", "Set to 1 to force FakeLLM mode — no LLM API calls, deterministic output."),
    ],
)

# ==============================================================
# 8. EVALUATION CRITERIA ALIGNMENT
# ==============================================================
add_heading("8. Evaluation Criteria Alignment")
add_table(
    ["Criterion", "Weight", "How it is met"],
    [
        ("Accuracy & completeness", "40%", "Facts carry source attribution; synthesis is strictly fact-bound; batching ensures all files are processed."),
        ("No hallucinations", "20%", "temperature=0; extract/synthesize separation; prompts forbid invention; empty fields instead of guesses; coerce_document enforces schema."),
        ("Token efficiency", "10%", "Files truncated at 20,000 chars; single extraction call per batch; compact fact JSON passed to synthesis."),
        ("LLM call count", "10%", "Typically 2 calls (1 extraction + 1 synthesis); fact caching eliminates re-extraction of unchanged files."),
        ("Execution time", "10%", "Per-step timing in trace; total duration in metadata.duration_ms."),
        ("Structure & format", "10%", "Pydantic response schema + coerce_document normalization guarantees well-formed output every time."),
    ],
)

# ==============================================================
# 9. SAMPLE DATA
# ==============================================================
add_heading("9. Sample Data")
add_para(
    "The sample_data/ directory contains four files that deliberately include a timeline conflict "
    "(8 weeks in project_brief.md vs 10 weeks in meeting_notes.md) and a budget discrepancy "
    "($150,000 in budget.json vs $120,000 in meeting_notes.md). These verify that the pipeline "
    "detects and surfaces conflicts in the trace rather than silently choosing one value."
)
add_table(
    ["File", "Contents"],
    [
        (("project_brief.md", True), "Project overview, 3 goals, timeline (8 weeks from June 1 2026), 3 team members."),
        (("meeting_notes.md", True), "Updated goals, conflicting timeline (10 weeks), alternative budget ($120K), 1 additional team member."),
        (("engineering_notes.txt", True), "Technical solution (microservices: FastAPI, PostgreSQL, Redis, RabbitMQ), architecture components, 3 identified risks."),
        (("budget.json", True), "Budget breakdown: total $150,000 USD (infrastructure $30K, licensing $15K, labor $105K)."),
    ],
)

# ==============================================================
# 10. PROJECT STRUCTURE
# ==============================================================
add_heading("10. Project Structure")
for line in [
    "Task 1/",
    "  app/",
    "    __init__.py",
    "    main.py           # FastAPI app & routes",
    "    pipeline.py       # Two-stage pipeline orchestration",
    "    ingest.py         # File normalization & batching prep",
    "    llm.py            # LLM backends + FakeLLM",
    "    prompts.py        # Extraction & synthesis prompts",
    "    schema.py         # Pydantic schemas + coerce_document()",
    "    jsonutil.py       # Robust JSON extraction from LLM output",
    "    dashboard.html    # Browser UI",
    "  sample_data/        # Example project files (with intentional conflicts)",
    "  tests/",
    "    test_pipeline.py",
    "    test_api_client.py",
    "  config.json         # Persistent API key / model configuration",
    "  requirements.txt",
    "  run_cli.py          # CLI entry point",
    "  start_server.bat    # Windows: start FastAPI server",
    "  run_cli_simulation.bat  # Windows: offline CLI run over sample_data",
    "  run_cli_live.bat        # Windows: live CLI run over any directory",
]:
    add_code(line)
doc.add_paragraph()

# ==============================================================
# 11. DEPENDENCIES
# ==============================================================
add_heading("11. Dependencies")
add_table(
    ["Package", "Purpose"],
    [
        (("fastapi >= 0.110", True), "HTTP API framework"),
        (("uvicorn[standard] >= 0.27", True), "ASGI server"),
        (("pydantic >= 2.5", True), "Request/response validation & schema enforcement"),
        (("anthropic >= 0.39", True), "Anthropic Claude client (optional; only used when AnthropicLLM is selected)"),
        (("requests", True), "HTTP calls for Gemini and OpenRouter backends (standard library not sufficient for JSON POST)"),
    ],
)

# ==============================================================
# SAVE
# ==============================================================
import os
out_path = os.path.join(os.path.dirname(__file__), "documentation.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
